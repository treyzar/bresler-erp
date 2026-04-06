import pytest
from decimal import Decimal
from io import BytesIO

from docx import Document

from apps.specs.services import document_service
from apps.specs.models import Specification
from .factories import (
    CommercialOfferFactory,
    OfferWorkItemFactory,
    SpecificationFactory,
    SpecificationLineFactory,
)


@pytest.mark.django_db
class TestGenerateOfferDocx:
    def test_generates_valid_docx(self):
        offer = CommercialOfferFactory()
        SpecificationFactory(offer=offer)

        buf = document_service.generate_offer_docx(offer)

        assert isinstance(buf, BytesIO)
        doc = Document(buf)
        assert len(doc.paragraphs) > 0

    def test_contains_offer_number(self):
        offer = CommercialOfferFactory(offer_number="4070/1-1")
        SpecificationFactory(offer=offer)

        buf = document_service.generate_offer_docx(offer)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "4070/1-1" in full_text

    def test_contains_participant_name(self):
        offer = CommercialOfferFactory()
        SpecificationFactory(offer=offer)
        participant_name = offer.participant.org_unit.name

        buf = document_service.generate_offer_docx(offer)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert participant_name in full_text

    def test_contains_spec_lines_in_table(self):
        offer = CommercialOfferFactory()
        spec = SpecificationFactory(offer=offer)
        SpecificationLineFactory(
            specification=spec, name="Терминал РЗА",
            quantity=2, unit_price=Decimal("50000"),
        )

        buf = document_service.generate_offer_docx(offer)
        doc = Document(buf)

        # Check table exists and has data row
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        row_texts = [cell.text for cell in table.rows[1].cells]
        assert "Терминал РЗА" in row_texts

    def test_contains_payment_terms(self):
        offer = CommercialOfferFactory(payment_terms="50_50")
        SpecificationFactory(offer=offer)

        buf = document_service.generate_offer_docx(offer)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "50%" in full_text

    def test_contains_works(self):
        offer = CommercialOfferFactory()
        SpecificationFactory(offer=offer)
        wi = OfferWorkItemFactory(offer=offer, included=True, days=10, specialists=2, trips=1)

        buf = document_service.generate_offer_docx(offer)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert wi.work_type.name in full_text
        assert "10 дн." in full_text
        assert "2 спец." in full_text

    def test_contains_additional_conditions(self):
        offer = CommercialOfferFactory(additional_conditions="Соответствие МЭК 61850")
        SpecificationFactory(offer=offer)

        buf = document_service.generate_offer_docx(offer)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "МЭК 61850" in full_text

    def test_no_additional_conditions_when_empty(self):
        offer = CommercialOfferFactory(additional_conditions="")
        SpecificationFactory(offer=offer)

        buf = document_service.generate_offer_docx(offer)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Дополнительные условия" not in full_text

    def test_delivery_included(self):
        offer = CommercialOfferFactory(delivery_included=True, delivery_city="Челябинск")
        SpecificationFactory(offer=offer)

        buf = document_service.generate_offer_docx(offer)
        doc = Document(buf)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "включена" in full_text
        assert "Челябинск" in full_text


@pytest.mark.django_db
class TestGenerateSpecificationDocx:
    def test_generates_valid_docx(self):
        offer = CommercialOfferFactory()
        spec = SpecificationFactory(offer=offer)
        SpecificationLineFactory(specification=spec, name="Бреслер-0107", quantity=1)

        buf = document_service.generate_specification_docx(offer)

        assert isinstance(buf, BytesIO)
        doc = Document(buf)
        assert len(doc.tables) >= 1

    def test_contains_lines(self):
        offer = CommercialOfferFactory()
        spec = SpecificationFactory(offer=offer)
        SpecificationLineFactory(specification=spec, name="Позиция А", quantity=3, unit_price=Decimal("1000"))
        SpecificationLineFactory(specification=spec, name="Позиция Б", quantity=1, unit_price=Decimal("5000"))

        buf = document_service.generate_specification_docx(offer)
        doc = Document(buf)

        table = doc.tables[0]
        all_cells_text = " ".join(cell.text for row in table.rows for cell in row.cells)
        assert "Позиция А" in all_cells_text
        assert "Позиция Б" in all_cells_text


@pytest.mark.django_db
class TestExportEndpoints:
    def test_export_offer_docx(self, authenticated_client):
        offer = CommercialOfferFactory()
        SpecificationFactory(offer=offer)

        resp = authenticated_client.get(f"/api/offers/{offer.pk}/export/")

        assert resp.status_code == 200
        assert "application/vnd.openxmlformats" in resp["Content-Type"]
        assert "attachment" in resp["Content-Disposition"]

    def test_export_specification_docx(self, authenticated_client):
        offer = CommercialOfferFactory()
        spec = SpecificationFactory(offer=offer)
        SpecificationLineFactory(specification=spec, name="Test", quantity=1)

        resp = authenticated_client.get(f"/api/offers/{offer.pk}/specification/export/")

        assert resp.status_code == 200
        assert "application/vnd.openxmlformats" in resp["Content-Type"]
