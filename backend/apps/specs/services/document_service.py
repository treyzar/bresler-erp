"""DOCX generation for CommercialOffer and Specification."""

import io
from decimal import Decimal

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from apps.specs.models import CommercialOffer, Specification


def _set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold


def _format_money(value) -> str:
    if value is None:
        return "—"
    return f"{value:,.2f}".replace(",", " ")


def _add_spec_table(doc, lines, vat_rate):
    """Add specification table to document."""
    headers = ["№", "Наименование", "Кол-во", "Цена за ед., руб.", "Итого, руб."]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    total = Decimal("0")
    for idx, line in enumerate(lines, 1):
        row = table.add_row().cells
        _set_cell_text(row[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row[1], line.name)
        _set_cell_text(row[2], str(line.quantity), align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row[3], _format_money(line.unit_price), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(row[4], _format_money(line.total_price), align=WD_ALIGN_PARAGRAPH.RIGHT)
        total += line.total_price

    # Totals row
    vat_amount = (total * vat_rate / 100).quantize(Decimal("0.01"))
    total_with_vat = total + vat_amount

    row = table.add_row().cells
    _set_cell_text(row[0], "")
    _set_cell_text(row[1], "")
    _set_cell_text(row[2], "")
    _set_cell_text(row[3], "Итого без НДС:", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_text(row[4], _format_money(total), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    row = table.add_row().cells
    _set_cell_text(row[3], f"НДС ({vat_rate}%):", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_text(row[4], _format_money(vat_amount), align=WD_ALIGN_PARAGRAPH.RIGHT)

    row = table.add_row().cells
    _set_cell_text(row[3], "Итого с НДС:", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_text(row[4], _format_money(total_with_vat), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    return total, vat_amount, total_with_vat


def _payment_terms_text(offer: CommercialOffer) -> str:
    """Generate human-readable payment terms."""
    templates = {
        "50_50": "50% предварительная оплата, 50% перед отгрузкой",
        "100_post_7": "100% в течение 7 рабочих дней после отгрузки",
        "100_post_30": "100% в течение 30 дней после отгрузки",
    }
    if offer.payment_terms in templates:
        return templates[offer.payment_terms]
    parts = []
    if offer.advance_percent:
        parts.append(f"{offer.advance_percent}% аванс")
    if offer.pre_shipment_percent:
        parts.append(f"{offer.pre_shipment_percent}% перед отгрузкой")
    if offer.post_payment_percent:
        parts.append(f"{offer.post_payment_percent}% постоплата")
    return ", ".join(parts) if parts else "по согласованию"


def _styled_paragraph(doc, text="", bold=False, alignment=None, size=12):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p, run


def generate_offer_docx(offer: CommercialOffer) -> io.BytesIO:
    """Generate complete КП as DOCX, returns BytesIO buffer."""
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Header ──
    _styled_paragraph(doc, offer.participant.org_unit.name,
                      bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()
    _styled_paragraph(doc, "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ",
                      bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    _styled_paragraph(doc, f"{offer.offer_number} от {offer.date:%d.%m.%Y}",
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ── Intro ──
    doc.add_paragraph(
        "ООО НПП «Бреслер» выражает благодарность за проявленный интерес к нашей продукции "
        "и предлагает к поставке следующее оборудование:"
    )
    doc.add_paragraph()

    # ── Specification table ──
    try:
        spec = offer.specification
        lines = list(spec.lines.all().order_by("line_number"))
    except Specification.DoesNotExist:
        lines = []

    if lines:
        _add_spec_table(doc, lines, offer.vat_rate)
    else:
        doc.add_paragraph("(спецификация не заполнена)")

    doc.add_paragraph()

    # ── Conditions ──
    _styled_paragraph(doc, "Условия:", bold=True)

    conditions = []
    conditions.append(f"1. Условия оплаты: {_payment_terms_text(offer)}.")
    conditions.append(f"2. {offer.shipment_condition_text}")
    conditions.append(f"3. Срок изготовления: {offer.manufacturing_period} дней.")
    conditions.append(f"4. Гарантийный срок: {offer.warranty_months} месяцев.")

    delivery = "включена" if offer.delivery_included else "не включена"
    if offer.delivery_included and offer.delivery_city:
        delivery += f" (г. {offer.delivery_city})"
    conditions.append(f"5. Доставка: {delivery}.")

    for c in conditions:
        doc.add_paragraph(c)

    # ── Works ──
    included_works = offer.work_items.filter(included=True).select_related("work_type")
    if included_works.exists():
        doc.add_paragraph()
        _styled_paragraph(doc, "Работы:", bold=True)
        for wi in included_works:
            doc.add_paragraph(
                f"— {wi.work_type.name}: {wi.days} дн., "
                f"{wi.specialists} спец., {wi.trips} выезд(ов)"
            )

    # ── Additional conditions ──
    if offer.additional_conditions:
        doc.add_paragraph()
        _styled_paragraph(doc, "Дополнительные условия:", bold=True)
        doc.add_paragraph(offer.additional_conditions)

    # ── Validity & contact ──
    doc.add_paragraph()
    valid_until = offer.valid_until.strftime("%d.%m.%Y") if offer.valid_until else "—"
    doc.add_paragraph(
        f"Срок действия предложения: {offer.valid_days} дней (до {valid_until})."
    )
    doc.add_paragraph()
    doc.add_paragraph("По вопросам обращаться:")
    manager_name = offer.manager.get_full_name() if offer.manager else "—"
    doc.add_paragraph(manager_name)

    # ── Save to buffer ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_specification_docx(offer: CommercialOffer) -> io.BytesIO:
    """Generate standalone specification as DOCX, returns BytesIO buffer."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Header
    _styled_paragraph(doc, "СПЕЦИФИКАЦИЯ",
                      bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    _styled_paragraph(doc, f"к коммерческому предложению {offer.offer_number} от {offer.date:%d.%m.%Y}",
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph(f"Заказчик: {offer.participant.org_unit.name}")
    doc.add_paragraph()

    # Table
    try:
        spec = offer.specification
        lines = list(spec.lines.all().order_by("line_number"))
    except Specification.DoesNotExist:
        lines = []

    if lines:
        _add_spec_table(doc, lines, offer.vat_rate)
    else:
        doc.add_paragraph("(спецификация не заполнена)")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
