import io
import base64
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os

class DocxRenderer:
    def __init__(self, template_path=None):
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
        else:
            self.doc = Document()
            # Установка стилей по умолчанию, если шаблона нет
            style = self.doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

    def render(self, json_blocks: list):
        # Сортируем блоки по Y, затем по X (на всякий случай, если фронт прислал не отсортированные)
        sorted_blocks = sorted(json_blocks, key=lambda b: (b.get('position', {}).get('y', 0), b.get('position', {}).get('x', 0)))
        
        last_y = 0
        pixels_per_inch = 96 # Стандарт для веба и A4 расчетов

        for block in sorted_blocks:
            block_type = block.get('type')
            content = block.get('content')
            style_cfg = block.get('style', {})
            pos = block.get('position', {})
            
            x_px = pos.get('x', 0)
            y_px = pos.get('y', 0)
            width_px = pos.get('width', 300)

            # Добавляем вертикальный отступ через пустые строки (приблизительно)
            # 12pt шрифт примерно 16px высотой
            if y_px > last_y + 10:
                diff = y_px - last_y
                empty_lines = int(diff / 20) # Приблизительно одна строка на 20px
                for _ in range(max(0, empty_lines)):
                    self.doc.add_paragraph()
            
            # Рассчитываем отступ слева в дюймах
            left_indent = Inches(x_px / pixels_per_inch)

            if block_type == 'paragraph':
                self._add_paragraph(content, style_cfg, left_indent)
            elif block_type == 'heading':
                self._add_heading(content, style_cfg, left_indent)
            elif block_type == 'table':
                # Для таблиц отступ делается сложнее, пока просто добавляем
                self._add_table(content, style_cfg)
            elif block_type == 'image':
                # Для изображений используем ширину из pos
                width_inches = Inches(width_px / pixels_per_inch)
                self._add_image(content, style_cfg, left_indent, width_inches)
            
            # Обновляем last_y (приблизительно высота элемента)
            # Текстовые блоки могут быть многострочными, это сложно учесть точно
            height_px = pos.get('height', 40)
            last_y = y_px + height_px
        
        return self.save_to_buffer()

    def _add_paragraph(self, content, style_cfg, left_indent=None):
        p = self.doc.add_paragraph(content)
        if left_indent:
            p.paragraph_format.left_indent = left_indent
        self._apply_style(p, style_cfg)

    def _add_heading(self, content, style_cfg, left_indent=None):
        h = self.doc.add_heading(content, level=1)
        if left_indent:
            h.paragraph_format.left_indent = left_indent
        self._apply_style(h, style_cfg)

    def _add_table(self, content, style_cfg):
        # content ожидается как список списков (rows -> cells)
        if not content or not isinstance(content, list):
            return

        try:
            rows_count = len(content)
            cols_count = len(content[0]) if rows_count > 0 and isinstance(content[0], list) else 0
            
            if rows_count == 0 or cols_count == 0:
                return

            table = self.doc.add_table(rows=rows_count, cols=cols_count)
            table.autofit = True
            
            # Попытка установить стиль сетки, если он существует
            try:
                table.style = 'Table Grid'
            except:
                pass

            for r_idx, row_data in enumerate(content):
                if not isinstance(row_data, list):
                    continue
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx < cols_count:
                        table.cell(r_idx, c_idx).text = str(cell_text)
        except Exception as e:
            self.doc.add_paragraph(f"[Ошибка создания таблицы: {str(e)}]")

    def _add_image(self, content, style_cfg, left_indent=None, width_inches=None):
        try:
            if ',' in content:
                content = content.split(',')[1]
            
            image_data = base64.b64decode(content)
            image_stream = io.BytesIO(image_data)
            
            # В python-docx изображение добавляется в параграф
            p = self.doc.add_paragraph()
            if left_indent:
                p.paragraph_format.left_indent = left_indent
            
            run = p.add_run()
            # Используем ширину из свойств элемента или по умолчанию 6 дюймов
            run.add_picture(image_stream, width=width_inches or Inches(6))
        except Exception as e:
            self.doc.add_paragraph(f"[Ошибка вставки изображения: {str(e)}]")

    def _apply_style(self, element, style_cfg):
        if not style_cfg:
            return

        # Выравнивание
        align = style_cfg.get('align')
        if align:
            mapping = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            element.alignment = mapping.get(align, WD_ALIGN_PARAGRAPH.LEFT)

        # Жирный/Курсив (для параграфов и заголовков)
        if hasattr(element, 'runs') and len(element.runs) > 0:
            run = element.runs[0]
            if style_cfg.get('bold'):
                run.bold = True
            if style_cfg.get('italic'):
                run.italic = True

    def save_to_buffer(self):
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer
