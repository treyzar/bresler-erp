import io
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
import pdfplumber


def normalize_owner_id(request):
    if hasattr(request, 'user') and request.user and hasattr(request.user, 'id') and request.user.id:
        return request.user.id
    return 1


def extract_text_from_node(node):
    if isinstance(node, str):
        return node
    if isinstance(node, dict):
        if node.get('type') == 'text':
            return node.get('text', '')
        children = node.get('content', [])
        return ''.join(extract_text_from_node(child) for child in children)
    if isinstance(node, list):
        return ''.join(extract_text_from_node(child) for child in node)
    return ''


def editor_json_to_plain_text(content_json):
    if not content_json:
        return ''
    
    doc = content_json if isinstance(content_json, dict) else {}
    content = doc.get('content', [])
    
    lines = []
    for block in content:
        text = extract_text_from_node(block)
        if text:
            lines.append(text)
    
    return '\n'.join(lines)


def get_text_with_marks(node):
    if node.get('type') == 'text':
        text = node.get('text', '')
        marks = node.get('marks', [])
        
        for mark in marks:
            mark_type = mark.get('type', '')
            if mark_type == 'bold':
                text = f'<b>{text}</b>'
            elif mark_type == 'italic':
                text = f'<i>{text}</i>'
            elif mark_type == 'underline':
                text = f'<u>{text}</u>'
        
        return text
    return ''


def process_content_for_pdf(content):
    result = []
    for node in content:
        if node.get('type') == 'text':
            result.append(get_text_with_marks(node))
        elif 'content' in node:
            result.extend(process_content_for_pdf(node.get('content', [])))
    return result


def editor_json_to_docx_bytes(content_json):
    doc = Document()
    
    if not content_json:
        return io.BytesIO()
    
    content = content_json.get('content', [])
    
    for block in content:
        block_type = block.get('type', 'paragraph')
        block_content = block.get('content', [])
        
        if block_type == 'heading':
            level = block.get('attrs', {}).get('level', 1)
            p = doc.add_heading(level=min(level, 9))
            for node in block_content:
                add_run_to_paragraph(p, node)
        
        elif block_type == 'paragraph':
            p = doc.add_paragraph()
            for node in block_content:
                add_run_to_paragraph(p, node)
        
        elif block_type == 'bulletList':
            for item in block_content:
                if item.get('type') == 'listItem':
                    item_content = item.get('content', [])
                    for para in item_content:
                        if para.get('type') == 'paragraph':
                            p = doc.add_paragraph(style='List Bullet')
                            for node in para.get('content', []):
                                add_run_to_paragraph(p, node)
        
        elif block_type == 'orderedList':
            for item in block_content:
                if item.get('type') == 'listItem':
                    item_content = item.get('content', [])
                    for para in item_content:
                        if para.get('type') == 'paragraph':
                            p = doc.add_paragraph(style='List Number')
                            for node in para.get('content', []):
                                add_run_to_paragraph(p, node)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def add_run_to_paragraph(paragraph, node):
    if node.get('type') == 'text':
        text = node.get('text', '')
        marks = node.get('marks', [])
        
        run = paragraph.add_run(text)
        
        for mark in marks:
            mark_type = mark.get('type', '')
            if mark_type == 'bold':
                run.bold = True
            elif mark_type == 'italic':
                run.italic = True
            elif mark_type == 'underline':
                run.underline = True


def editor_json_to_pdf_bytes(content_json):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=72, rightMargin=72, topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='CustomBody', parent=styles['Normal'], fontSize=12, leading=16))
    styles.add(ParagraphStyle(name='CustomH1', parent=styles['Heading1'], fontSize=24, leading=28, spaceAfter=12))
    styles.add(ParagraphStyle(name='CustomH2', parent=styles['Heading2'], fontSize=20, leading=24, spaceAfter=10))
    styles.add(ParagraphStyle(name='CustomH3', parent=styles['Heading3'], fontSize=16, leading=20, spaceAfter=8))
    
    story = []
    
    if not content_json:
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    content = content_json.get('content', [])
    
    for block in content:
        block_type = block.get('type', 'paragraph')
        block_content = block.get('content', [])
        
        if block_type == 'heading':
            level = block.get('attrs', {}).get('level', 1)
            style_name = f'CustomH{min(level, 3)}'
            text_parts = process_content_for_pdf(block_content)
            text = ''.join(text_parts)
            if text:
                story.append(Paragraph(text, styles[style_name]))
                story.append(Spacer(1, 6))
        
        elif block_type == 'paragraph':
            text_parts = process_content_for_pdf(block_content)
            text = ''.join(text_parts)
            if text:
                story.append(Paragraph(text, styles['CustomBody']))
                story.append(Spacer(1, 4))
            else:
                story.append(Spacer(1, 12))
        
        elif block_type in ['bulletList', 'orderedList']:
            bullet_type = 'bullet' if block_type == 'bulletList' else '1'
            items = []
            for item in block_content:
                if item.get('type') == 'listItem':
                    item_content = item.get('content', [])
                    for para in item_content:
                        if para.get('type') == 'paragraph':
                            text_parts = process_content_for_pdf(para.get('content', []))
                            text = ''.join(text_parts)
                            if text:
                                items.append(ListItem(Paragraph(text, styles['CustomBody'])))
            
            if items:
                story.append(ListFlowable(items, bulletType=bullet_type))
                story.append(Spacer(1, 8))
    
    if not story:
        story.append(Spacer(1, 1))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def docx_file_to_editor_json(file):
    doc = Document(file)
    
    content = []
    
    for para in doc.paragraphs:
        if not para.text.strip():
            content.append({
                'type': 'paragraph',
                'content': []
            })
            continue
        
        style_name = para.style.name.lower() if para.style else ''
        
        if 'heading' in style_name:
            level = 1
            for i in range(1, 10):
                if str(i) in style_name:
                    level = i
                    break
            
            block = {
                'type': 'heading',
                'attrs': {'level': level},
                'content': []
            }
        elif 'list bullet' in style_name:
            block = {
                'type': 'paragraph',
                'content': []
            }
        elif 'list number' in style_name:
            block = {
                'type': 'paragraph',
                'content': []
            }
        else:
            block = {
                'type': 'paragraph',
                'content': []
            }
        
        for run in para.runs:
            if run.text:
                text_node = {
                    'type': 'text',
                    'text': run.text
                }
                
                marks = []
                if run.bold:
                    marks.append({'type': 'bold'})
                if run.italic:
                    marks.append({'type': 'italic'})
                if run.underline:
                    marks.append({'type': 'underline'})
                
                if marks:
                    text_node['marks'] = marks
                
                block['content'].append(text_node)
        
        if block['content'] or block['type'] == 'paragraph':
            content.append(block)
    
    return {
        'type': 'doc',
        'content': content
    }


def pdf_file_to_editor_json(file):
    content = []
    
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                paragraphs = text.split('\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        content.append({
                            'type': 'paragraph',
                            'content': [
                                {
                                    'type': 'text',
                                    'text': para_text
                                }
                            ]
                        })
                    else:
                        content.append({
                            'type': 'paragraph',
                            'content': []
                        })
    
    return {
        'type': 'doc',
        'content': content
    }
