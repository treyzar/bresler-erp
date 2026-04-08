"""
Tests for Tiptap JSON → DOCX/PDF conversion
Ensures proper handling of all Tiptap node types
"""
import io
import unittest
import sys
import os

# Add backend to path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..', '..', '..'))

from apps.edo.doc_builder.services.converters import (
    editor_json_to_docx_bytes,
    editor_json_to_pdf_bytes,
    editor_json_to_plain_text,
)
from docx import Document


class TestTiptapToDocxConversion(unittest.TestCase):
    """Test DOCX conversion from Tiptap JSON"""

    def test_convert_simple_paragraph(self):
        """Test conversion of simple paragraph"""
        content_json = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [
                        {"type": "text", "text": "Hello World"}
                    ]
                }
            ]
        }

        result = editor_json_to_docx_bytes(content_json)
        doc = Document(io.BytesIO(result))

        self.assertEqual(len(doc.paragraphs), 1)
        self.assertEqual(doc.paragraphs[0].text, "Hello World")

    def test_convert_paragraph_with_marks(self):
        """Test conversion with bold, italic, underline"""
        content_json = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [
                        {
                            "type": "text",
                            "text": "Bold Text",
                            "marks": [{"type": "bold"}]
                        }
                    ]
                }
            ]
        }

        result = editor_json_to_docx_bytes(content_json)
        doc = Document(io.BytesIO(result))

        self.assertEqual(doc.paragraphs[0].text, "Bold Text")
        self.assertTrue(doc.paragraphs[0].runs[0].bold)

    def test_convert_textstyle_marks(self):
        """Test conversion with textStyle marks (color, fontSize, fontFamily)"""
        content_json = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [
                        {
                            "type": "text",
                            "text": "Styled Text",
                            "marks": [
                                {
                                    "type": "textStyle",
                                    "attrs": {
                                        "color": "#FF0000",
                                        "fontSize": "16px",
                                        "fontFamily": "Arial"
                                    }
                                }
                            ]
                        }
                    ]
                }
            ]
        }

        result = editor_json_to_docx_bytes(content_json)
        doc = Document(io.BytesIO(result))

        run = doc.paragraphs[0].runs[0]
        self.assertEqual(run.text, "Styled Text")

    def test_convert_heading(self):
        """Test conversion of headings"""
        content_json = {
            "type": "doc",
            "content": [
                {
                    "type": "heading",
                    "attrs": {"level": 1},
                    "content": [
                        {"type": "text", "text": "Heading 1"}
                    ]
                }
            ]
        }

        result = editor_json_to_docx_bytes(content_json)
        doc = Document(io.BytesIO(result))

        self.assertEqual(doc.paragraphs[0].text, "Heading 1")
        self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")

    def test_convert_simple_table(self):
        """Test conversion of simple 2x2 table"""
        content_json = {
            "type": "doc",
            "content": [
                {
                    "type": "table",
                    "content": [
                        {
                            "type": "tableRow",
                            "content": [
                                {
                                    "type": "tableCell",
                                    "content": [
                                        {
                                            "type": "paragraph",
                                            "content": [{"type": "text", "text": "A1"}]
                                        }
                                    ]
                                },
                                {
                                    "type": "tableCell",
                                    "content": [
                                        {
                                            "type": "paragraph",
                                            "content": [{"type": "text", "text": "B1"}]
                                        }
                                    ]
                                }
                            ]
                        },
                        {
                            "type": "tableRow",
                            "content": [
                                {
                                    "type": "tableCell",
                                    "content": [
                                        {
                                            "type": "paragraph",
                                            "content": [{"type": "text", "text": "A2"}]
                                        }
                                    ]
                                },
                                {
                                    "type": "tableCell",
                                    "content": [
                                        {
                                            "type": "paragraph",
                                            "content": [{"type": "text", "text": "B2"}]
                                        }
                                    ]
                                }
                            ]
                        }
                    ]
                }
            ]
        }

        result = editor_json_to_docx_bytes(content_json)
        doc = Document(io.BytesIO(result))

        self.assertEqual(len(doc.tables), 1)
        table = doc.tables[0]
        self.assertEqual(len(table.rows), 2)
        self.assertEqual(len(table.columns), 2)
        self.assertEqual(table.rows[0].cells[0].text, "A1")
        self.assertEqual(table.rows[1].cells[1].text, "B2")

    def test_convert_table_with_colspan(self):
        """Test conversion of table with colspan"""
        content_json = {
            "type": "doc",
            "content": [
                {
                    "type": "table",
                    "content": [
                        {
                            "type": "tableRow",
                            "content": [
                                {
                                    "type": "tableCell",
                                    "attrs": {"colspan": 2, "rowspan": 1},
                                    "content": [
                                        {
                                            "type": "paragraph",
                                            "content": [{"type": "text", "text": "Merged"}]
                                        }
                                    ]
                                }
                            ]
                        },
                        {
                            "type": "tableRow",
                            "content": [
                                {
                                    "type": "tableCell",
                                    "content": [
                                        {
                                            "type": "paragraph",
                                            "content": [{"type": "text", "text": "A2"}]
                                        }
                                    ]
                                },
                                {
                                    "type": "tableCell",
                                    "content": [
                                        {
                                            "type": "paragraph",
                                            "content": [{"type": "text", "text": "B2"}]
                                        }
                                    ]
                                }
                            ]
                        }
                    ]
                }
            ]
        }

        result = editor_json_to_docx_bytes(content_json)
        doc = Document(io.BytesIO(result))

        table = doc.tables[0]
        # First row should have 1 merged cell
        self.assertEqual(len(table.rows[0].cells), 1)
        self.assertEqual(table.rows[0].cells[0].text, "Merged")

    def test_convert_bullet_list(self):
        """Test conversion of bullet list"""
        content_json = {
            "type": "doc",
            "content": [
                {
                    "type": "bulletList",
                    "content": [
                        {
                            "type": "listItem",
                            "content": [
                                {
                                    "type": "paragraph",
                                    "content": [{"type": "text", "text": "Item 1"}]
                                }
                            ]
                        },
                        {
                            "type": "listItem",
                            "content": [
                                {
                                    "type": "paragraph",
                                    "content": [{"type": "text", "text": "Item 2"}]
                                }
                            ]
                        }
                    ]
                }
            ]
        }

        result = editor_json_to_docx_bytes(content_json)
        doc = Document(io.BytesIO(result))

        # Should have 2 paragraphs with List Bullet style
        bullet_paragraphs = [
            p for p in doc.paragraphs 
            if p.style.name == "List Bullet"
        ]
        self.assertEqual(len(bullet_paragraphs), 2)

    def test_convert_ordered_list(self):
        """Test conversion of ordered list"""
        content_json = {
            "type": "doc",
            "content": [
                {
                    "type": "orderedList",
                    "content": [
                        {
                            "type": "listItem",
                            "content": [
                                {
                                    "type": "paragraph",
                                    "content": [{"type": "text", "text": "First"}]
                                }
                            ]
                        }
                    ]
                }
            ]
        }

        result = editor_json_to_docx_bytes(content_json)
        doc = Document(io.BytesIO(result))

        numbered_paragraphs = [
            p for p in doc.paragraphs 
            if p.style.name == "List Number"
        ]
        self.assertEqual(len(numbered_paragraphs), 1)

    def test_convert_empty_document(self):
        """Test conversion of empty document"""
        content_json = {
            "type": "doc",
            "content": []
        }

        result = editor_json_to_docx_bytes(content_json)
        doc = Document(io.BytesIO(result))

        self.assertEqual(len(doc.paragraphs), 0)

    def test_legacy_canvas_format(self):
        """Test backward compatibility with old Canvas format"""
        canvas_elements = [
            {
                "type": "text",
                "y": 100,
                "x": 50,
                "properties": {
                    "content": "Legacy Text",
                    "bold": True,
                }
            }
        ]

        result = editor_json_to_docx_bytes(canvas_elements)
        doc = Document(io.BytesIO(result))

        self.assertEqual(len(doc.paragraphs), 1)
        self.assertEqual(doc.paragraphs[0].text, "Legacy Text")


class TestTiptapToPdfConversion(unittest.TestCase):
    """Test PDF conversion from Tiptap JSON"""

    def test_convert_simple_paragraph_pdf(self):
        """Test PDF conversion of simple paragraph"""
        content_json = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [
                        {"type": "text", "text": "Hello PDF"}
                    ]
                }
            ]
        }

        result = editor_json_to_pdf_bytes(content_json)
        self.assertIsInstance(result, bytes)
        self.assertGreater(len(result), 0)

    def test_convert_table_pdf(self):
        """Test PDF conversion of table"""
        content_json = {
            "type": "doc",
            "content": [
                {
                    "type": "table",
                    "content": [
                        {
                            "type": "tableRow",
                            "content": [
                                {
                                    "type": "tableCell",
                                    "content": [
                                        {
                                            "type": "paragraph",
                                            "content": [{"type": "text", "text": "Cell"}]
                                        }
                                    ]
                                }
                            ]
                        }
                    ]
                }
            ]
        }

        result = editor_json_to_pdf_bytes(content_json)
        self.assertIsInstance(result, bytes)
        self.assertGreater(len(result), 0)

    def test_convert_empty_pdf(self):
        """Test PDF conversion of empty document"""
        content_json = {
            "type": "doc",
            "content": []
        }

        result = editor_json_to_pdf_bytes(content_json)
        self.assertIsInstance(result, bytes)


class TestPlainExtraction(unittest.TestCase):
    """Test plain text extraction"""

    def test_extract_plain_text(self):
        """Test extracting plain text from Tiptap JSON"""
        content_json = {
            "type": "doc",
            "content": [
                {
                    "type": "paragraph",
                    "content": [
                        {"type": "text", "text": "Line 1"}
                    ]
                },
                {
                    "type": "paragraph",
                    "content": [
                        {"type": "text", "text": "Line 2"}
                    ]
                }
            ]
        }

        result = editor_json_to_plain_text(content_json)
        self.assertEqual(result, "Line 1\nLine 2")

    def test_extract_empty_text(self):
        """Test extracting from empty document"""
        content_json = {
            "type": "doc",
            "content": []
        }

        result = editor_json_to_plain_text(content_json)
        self.assertEqual(result, "")


if __name__ == "__main__":
    unittest.main()
