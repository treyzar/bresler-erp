import uuid
import re
from django.db import models
from django.utils import timezone

class Template(models.Model):
    TEMPLATE_TYPE_CHOICES = [
        ('HTML', 'HTML'),
        ('DOCX', 'DOCX'),
        ('PDF', 'PDF'),  # Добавили PDF как тип
    ]

    VISIBILITY_CHOICES = [
        ('PUBLIC', 'Public'),
        ('RESTRICTED', 'Restricted'),
    ]

    title = models.CharField(max_length=255)
    description = models.TextField(blank=True, default='')
    template_type = models.CharField(max_length=10, choices=TEMPLATE_TYPE_CHOICES)
    visibility = models.CharField(max_length=20, choices=VISIBILITY_CHOICES, default='PUBLIC')
    owner_id = models.IntegerField(default=1)
    allowed_users = models.JSONField(default=list, blank=True)
    is_deleted = models.BooleanField(default=False, db_index=True)
    
    # HTML нужен для генерации PDF (рендера)
    html_content = models.TextField(blank=True, default='')
    
    # НОВОЕ ПОЛЕ: Хранит "сырое" состояние редактора (координаты, типы элементов и т.д.)
    editor_content = models.JSONField(default=list, blank=True)
    
    docx_file = models.FileField(upload_to='docx_templates/', blank=True, null=True)
    
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    def save(self, *args, **kwargs):
        # html_content — производный кеш от editor_content.
        # Регенерируем, если editor_content непустой. Для DOCX-импорта, где есть
        # готовый html_content, но нет editor_content, оставляем как пришло.
        if self.editor_content:
            from ..services.html_renderer import render_editor_content_html
            self.html_content = render_editor_content_html(self.editor_content)
        super().save(*args, **kwargs)

    def get_placeholders(self):
        # Логика для HTML/PDF шаблонов, которые рендерятся через HTML
        if self.template_type in ['HTML', 'PDF']:
            pattern = r'\{\{\s*(\w+)\s*\}\}'
            sources = [self.html_content or '']
            # Заодно вычитываем плейсхолдеры из editor_content, чтобы не зависеть
            # от того, как именно отрендерен HTML (escape, CSS, и т.п.).
            for el in (self.editor_content or []):
                if not isinstance(el, dict):
                    continue
                props = el.get('properties') or {}
                for key in ('content', 'text'):
                    val = props.get(key)
                    if isinstance(val, str):
                        sources.append(val)
                cells = props.get('cells')
                if isinstance(cells, list):
                    for row in cells:
                        if isinstance(row, list):
                            for cell in row:
                                if isinstance(cell, dict):
                                    val = cell.get('content')
                                    if isinstance(val, str):
                                        sources.append(val)
                data = props.get('data')
                if isinstance(data, list):
                    for row in data:
                        if isinstance(row, list):
                            for cell in row:
                                if isinstance(cell, str):
                                    sources.append(cell)
            found: set[str] = set()
            for src in sources:
                found.update(re.findall(pattern, src))
            return sorted(found)

        elif self.template_type == 'DOCX' and self.docx_file:
            from docx import Document
            try:
                doc = Document(self.docx_file.path)
                text = '\n'.join([p.text for p in doc.paragraphs])
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += '\n' + cell.text
                pattern = r'\{\{\s*(\w+)\s*\}\}'
                return list(set(re.findall(pattern, text)))
            except Exception:
                return []
        return []

    def is_accessible_by(self, user_id):
        if self.visibility == 'PUBLIC':
            return True
        if self.owner_id == user_id:
            return True
        if user_id in self.allowed_users:
            return True
        return False

    def __str__(self):
        return self.title

class TemplateVersion(models.Model):
    template = models.ForeignKey(Template, on_delete=models.CASCADE, related_name='versions')
    version_number = models.IntegerField()
    
    html_content = models.TextField(blank=True, default='')
    # Также сохраняем JSON структуру в версиях
    editor_content = models.JSONField(default=list, blank=True)
    
    docx_file = models.FileField(upload_to='docx_versions/', blank=True, null=True)
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        ordering = ['-version_number']
        unique_together = ['template', 'version_number']

    def __str__(self):
        return f"{self.template.title} v{self.version_number}"

class ShareLink(models.Model):
    template = models.ForeignKey(Template, on_delete=models.CASCADE, related_name='share_links')
    token = models.CharField(max_length=64, unique=True, default=uuid.uuid4)
    ttl_days = models.IntegerField(default=7)
    max_uses = models.IntegerField(default=50)
    current_uses = models.IntegerField(default=0)
    created_at = models.DateTimeField(auto_now_add=True)

    def is_valid(self):
        if self.current_uses >= self.max_uses:
            return False
        expiry = self.created_at + timezone.timedelta(days=self.ttl_days)
        if timezone.now() > expiry:
            return False
        return True

    def increment_use(self):
        self.current_uses += 1
        self.save()

    def __str__(self):
        return f"ShareLink for {self.template.title}"
