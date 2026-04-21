import base64
import io
import logging
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .docx_table_builder import DocxTableBuilder

logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{\s*(\w+)\s*\}\}")

# 96 DPI — ровно то, чем живёт холст редактора; переводим пиксели в дюймы для python-docx.
PX_TO_INCHES = 1.0 / 96.0


def _apply_placeholders(text: str, values: dict) -> str:
    if not text or not values:
        return text or ""

    def _sub(match: re.Match) -> str:
        key = match.group(1)
        val = values.get(key)
        return "" if val is None else str(val)

    return _PLACEHOLDER_RE.sub(_sub, text)


def _props(el: dict) -> dict:
    props = el.get("properties")
    return props if isinstance(props, dict) else {}


def _hex_to_rgb(color: str | None) -> RGBColor | None:
    if not color or not isinstance(color, str):
        return None
    c = color.strip().lstrip("#")
    if len(c) == 3:
        c = "".join(ch * 2 for ch in c)
    if len(c) != 6:
        return None
    try:
        return RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))
    except ValueError:
        return None


def _add_text_run(paragraph, text: str, props: dict) -> None:
    run = paragraph.add_run(text)
    if props.get("bold"):
        run.bold = True
    if props.get("italic"):
        run.italic = True
    if props.get("underline"):
        run.underline = True
    font_size = props.get("fontSize")
    if font_size:
        try:
            run.font.size = Pt(float(font_size))
        except (TypeError, ValueError):
            pass
    font_family = props.get("fontFamily")
    if font_family:
        run.font.name = str(font_family)
    rgb = _hex_to_rgb(props.get("color"))
    if rgb is not None:
        run.font.color.rgb = rgb


def _add_image_run(paragraph, data_url: str, width_px: float | int, height_px: float | int) -> bool:
    """Вставляет data-URL картинку как inline-image. Возвращает True при успехе."""
    if not data_url or not data_url.startswith("data:image/"):
        return False
    try:
        _, payload = data_url.split(",", 1)
    except ValueError:
        return False
    try:
        img_bytes = base64.b64decode(payload)
    except Exception:
        return False
    stream = io.BytesIO(img_bytes)
    try:
        run = paragraph.add_run()
        width_in = max(0.1, float(width_px) * PX_TO_INCHES)
        height_in = max(0.1, float(height_px) * PX_TO_INCHES)
        run.add_picture(stream, width=Inches(width_in), height=Inches(height_in))
        return True
    except Exception as e:
        logger.warning("Failed to embed image in DOCX: %s", e)
        return False


class DocxBuilderService:
    Y_TOLERANCE = 15  # px

    @classmethod
    def build_from_json(cls, editor_content: list, values: dict | None = None) -> bytes:
        """Собирает DOCX из editor_content с группировкой в строки по Y.

        Ожидает каноничный формат элементов с `properties`. Плейсхолдеры
        {{var}} подставляются на уровне содержимого, не в готовом DOCX.
        """
        doc = Document()
        values = values or {}

        if not editor_content:
            logger.warning("DocxBuilder: No editor content provided.")
            return b""

        try:
            sorted_by_y = sorted(editor_content, key=lambda e: float(e.get("y", 0) or 0))

            rows: list[list[dict]] = []
            if sorted_by_y:
                current_row = [sorted_by_y[0]]
                current_base_y = float(sorted_by_y[0].get("y", 0) or 0)
                for el in sorted_by_y[1:]:
                    el_y = float(el.get("y", 0) or 0)
                    if abs(el_y - current_base_y) <= cls.Y_TOLERANCE:
                        current_row.append(el)
                    else:
                        rows.append(current_row)
                        current_row = [el]
                        current_base_y = el_y
                if current_row:
                    rows.append(current_row)

            for row in rows:
                row.sort(key=lambda e: float(e.get("x", 0) or 0))

                # Таблицу рендерим как отдельный блок, без горизонтальной группировки.
                if len(row) == 1 and row[0].get("type") == "table":
                    DocxTableBuilder.render_table(doc, row[0], values)
                    continue

                p = doc.add_paragraph()

                if len(row) == 1:
                    align = _props(row[0]).get("align", "left")
                    if align == "center":
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif align == "right":
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                for idx, el in enumerate(row):
                    el_type = el.get("type")
                    props = _props(el)

                    if el_type == "text":
                        content = _apply_placeholders(str(props.get("content") or ""), values)
                        _add_text_run(p, content, props)

                    elif el_type == "signature":
                        image = props.get("image")
                        if image and _add_image_run(p, image, el.get("width") or 150, el.get("height") or 40):
                            pass
                        else:
                            text = _apply_placeholders(str(props.get("text") or "Подпись"), values)
                            _add_text_run(p, text, props)

                    elif el_type == "image":
                        src = props.get("src") or ""
                        if src and not _add_image_run(p, src, el.get("width") or 100, el.get("height") or 100):
                            _add_text_run(p, str(props.get("alt") or ""), props)

                    elif el_type == "date":
                        text = _apply_placeholders(str(props.get("value") or ""), values)
                        _add_text_run(p, text, props)

                    elif el_type == "divider":
                        _add_text_run(p, "\u2014" * 10, props)  # em-dashes

                    elif el_type == "table":
                        DocxTableBuilder.render_table(doc, el, values)

                    if el_type != "table" and idx < len(row) - 1:
                        next_x = float(row[idx + 1].get("x", 0) or 0)
                        this_right = float(el.get("x", 0) or 0) + float(el.get("width", 0) or 0)
                        gap = next_x - this_right
                        p.add_run("\t\t" if gap > 50 else "\t")

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()

        except Exception as e:
            logger.error("DocxBuilder Service Error: %s", e, exc_info=True)
            raise RuntimeError(f"Failed to build DOCX: {str(e)}")
