from __future__ import annotations

import io
import math
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from PIL import Image as PILImage

from .elements import (
    is_likely_signature,
    make_divider_element,
    make_image_element,
    make_signature_element,
    make_table_element,
    make_text_element,
)

# === ТОЧНЫЕ КОНСТАНТЫ WORD ===
PT_TO_PX = 1.3333333333333333  # 96 DPI / 72 DPI (pt to px)
MM_TO_PX = 3.937007874015748  # 96 DPI / 25.4 mm

# A4 при 96 DPI
DEFAULT_PAGE_WIDTH = int(210 * MM_TO_PX)  # 794px
DEFAULT_PAGE_HEIGHT = int(297 * MM_TO_PX)  # 1123px
# Для редактора используем непрерывные координаты по высоте страницы,
# без промежутков между страницами.
PAGE_GAP = 0

# Стандартные поля Word (2.54 см)
DEFAULT_MARGIN = int(25.4 * MM_TO_PX)  # 96px
DEFAULT_LEFT_MARGIN = DEFAULT_MARGIN
DEFAULT_RIGHT_MARGIN = DEFAULT_MARGIN
DEFAULT_TOP_MARGIN = DEFAULT_MARGIN
DEFAULT_BOTTOM_MARGIN = DEFAULT_MARGIN

PAGE_WIDTH = DEFAULT_PAGE_WIDTH
PAGE_HEIGHT = DEFAULT_PAGE_HEIGHT
LEFT_MARGIN = DEFAULT_LEFT_MARGIN
RIGHT_MARGIN = DEFAULT_RIGHT_MARGIN
TOP_MARGIN = DEFAULT_TOP_MARGIN
BOTTOM_MARGIN = DEFAULT_BOTTOM_MARGIN
CONTENT_WIDTH = PAGE_WIDTH - (LEFT_MARGIN + RIGHT_MARGIN)


def parse_docx(file_obj) -> dict[str, Any]:
    """Главная функция парсинга с АБСОЛЮТНЫМ позиционированием"""
    _reset_layout_defaults()
    doc = Document(file_obj)
    elements: list[dict[str, Any]] = []
    plain_chunks: list[str] = []
    y_offset = TOP_MARGIN
    current_page_idx = 0

    # Загрузка реальных полей документа
    _load_document_margins(doc)
    y_offset = TOP_MARGIN

    # Кэш стилей
    style_cache = _build_style_cache(doc)

    # Карта изображений
    images_map = {}
    try:
        for rel in doc.part.rels.values():
            if rel.target_ref and ("image" in rel.target_ref or "media" in rel.target_ref):
                images_map[rel.rId] = rel.target_part.blob
    except Exception:
        pass

    # === ОСНОВНОЙ ЦИКЛ ПАРСИНГА ===
    for kind, block, _section_idx in _iter_all_blocks(doc):
        if kind == "paragraph":
            # Разрыв страницы
            if _has_page_break(block):
                current_page_idx += 1
                y_offset = _page_top(current_page_idx)
                plain_chunks.append("[Page Break]")

            # Форматирование параграфа
            fmt = _get_paragraph_format(block, style_cache)

            # Отступ ПЕРЕД
            y_offset += fmt.space_before

            # Разделитель (граница)
            if fmt.has_border_bottom:
                elements.append(
                    make_divider_element(
                        x=LEFT_MARGIN, y=y_offset, width=CONTENT_WIDTH, thickness=max(1, fmt.border_width)
                    )
                )
                y_offset += max(2, fmt.border_width) + 6

            # === ИЗОБРАЖЕНИЯ В ПАРАГРАФЕ ===
            blips = block._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
            has_image = False
            for blip in blips:
                embed = blip.get(qn("r:embed"))
                if embed and embed in images_map:
                    has_image = True
                    # Изображения выравниваются по левому краю с отступом параграфа
                    img_x = LEFT_MARGIN + fmt.left_indent
                    el, h_added = _process_image(images_map[embed], img_x, y_offset)
                    if el:
                        elements.append(el)
                        y_offset += h_added + 10  # Отступ под изображением
                        plain_chunks.append("[Image]")

            # === ТЕКСТ ===
            text_content = _extract_paragraph_text(block)
            if text_content:
                # X с учетом всех отступов
                final_x = LEFT_MARGIN + fmt.left_indent + fmt.first_line_indent
                eff_width = max(50, CONTENT_WIDTH - fmt.left_indent - fmt.right_indent)

                # ТОЧНАЯ ВЫСОТА
                total_h = _calculate_text_height_exact(text_content, fmt.size, eff_width, fmt.line_spacing)

                elements.append(
                    make_text_element(
                        x=final_x,
                        y=y_offset,
                        width=eff_width,
                        height=total_h,
                        content=text_content,
                        font=fmt.font,
                        size=fmt.size,
                        bold=fmt.bold,
                        italic=fmt.italic,
                        color=fmt.color,
                        align=fmt.align,
                    )
                )
                y_offset += total_h
                plain_chunks.append(text_content)

            # Отступ ПОСЛЕ
            y_offset += fmt.space_after

            # Минимальный интервал для пустых параграфов
            if not text_content and not has_image and not fmt.has_border_bottom:
                y_offset += 8

            # Контроль высоты страницы
            if y_offset > _page_bottom_limit(current_page_idx):
                current_page_idx += 1
                y_offset = _page_top(current_page_idx)

        # === ОБРАБОТКА ТАБЛИЦЫ ===
        elif kind == "table":
            try:
                # Получаем форматирование таблицы (отступы, выравнивание)
                table_format = _get_table_format(block)

                # Добавляем отступ перед таблицей
                y_offset += table_format.get("space_before", 8)

                table_data = _parse_table(block, y_offset, style_cache)
                if table_data:
                    # Проверка, помещается ли таблица
                    if y_offset + table_data["total_height"] > _page_bottom_limit(current_page_idx, reserved=20):
                        current_page_idx += 1
                        y_offset = _page_top(current_page_idx)

                    # X позиция с учетом отступов таблицы
                    table_x = LEFT_MARGIN + table_format.get("left_indent", 0)
                    table_width = (
                        CONTENT_WIDTH - table_format.get("left_indent", 0) - table_format.get("right_indent", 0)
                    )

                    elements.append(
                        make_table_element(
                            x=table_x,
                            y=y_offset,
                            width=table_width,
                            height=table_data["total_height"],
                            data=table_data["rows"],
                            cell_text_colors=table_data.get("cell_text_colors"),
                        )
                    )
                    y_offset += table_data["total_height"] + table_format.get("space_after", 12)
                    plain_chunks.append("[Table]")
            except Exception as e:
                print(f"Table parsing error: {e}")
                import traceback

                traceback.print_exc()

    return {
        "elements": elements,
        "text": "\n".join(plain_chunks),
        "metadata": {
            "page_width": PAGE_WIDTH,
            "page_height": PAGE_HEIGHT,
            "page_gap": PAGE_GAP,
            "pages_estimated": current_page_idx + 1,
            "margins": {"left": LEFT_MARGIN, "right": RIGHT_MARGIN, "top": TOP_MARGIN, "bottom": BOTTOM_MARGIN},
        },
    }


def _build_style_cache(doc: DocumentObject) -> dict[str, ParagraphFormat]:
    """Кэширование всех стилей документа"""
    cache = {}
    try:
        for style in doc.styles:
            if style.type == 1:  # Paragraph style
                fmt = ParagraphFormat()

                style_fmt = style.paragraph_format
                if style_fmt:
                    if style_fmt.left_indent:
                        fmt.left_indent = int(style_fmt.left_indent.pt * PT_TO_PX)
                    if style_fmt.right_indent:
                        fmt.right_indent = int(style_fmt.right_indent.pt * PT_TO_PX)
                    if style_fmt.first_line_indent:
                        fmt.first_line_indent = int(style_fmt.first_line_indent.pt * PT_TO_PX)
                    if style_fmt.space_before:
                        fmt.space_before = int(style_fmt.space_before.pt * PT_TO_PX)
                    if style_fmt.space_after:
                        fmt.space_after = int(style_fmt.space_after.pt * PT_TO_PX)
                    if style_fmt.line_spacing:
                        fmt.line_spacing = int(style_fmt.line_spacing * PT_TO_PX) or int(fmt.size * 1.15)
                else:
                    fmt.line_spacing = int(fmt.size * 1.15)

                if style.font:
                    if style.font.bold:
                        fmt.bold = True
                    if style.font.italic:
                        fmt.italic = True
                    if style.font.size:
                        fmt.size = int(style.font.size.pt)
                    if style.font.name:
                        fmt.font = style.font.name

                cache[style.name] = fmt
    except Exception:
        pass

    if "Normal" not in cache:
        cache["Normal"] = ParagraphFormat()

    return cache


def _process_image(img_bytes, x, y):
    """Обработка изображения"""
    if not img_bytes:
        return None, 0

    try:
        pil_img = PILImage.open(io.BytesIO(img_bytes))
        w, h = pil_img.size

        max_width = CONTENT_WIDTH - (x - LEFT_MARGIN)
        scale = min(max_width / w, 1.0) if w > max_width else 1.0
        final_w = int(w * scale)
        final_h = int(h * scale)

        ext = (pil_img.format or "png").lower()

        if is_likely_signature(final_w, final_h):
            return make_signature_element(x, y, final_w, final_h, img_bytes, ext), final_h
        else:
            return make_image_element(x, y, final_w, final_h, img_bytes, ext), final_h
    except Exception:
        return None, 0


def _iter_all_blocks(doc: DocumentObject):
    """Обход всех блоков документа"""
    section_idx = 0
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield ("paragraph", Paragraph(child, doc), section_idx)
        elif child.tag.endswith("}tbl"):
            yield ("table", Table(child, doc), section_idx)
        elif child.tag.endswith("}AlternateContent"):
            for choice in child.iterchildren():
                for inner in choice.iterchildren():
                    if inner.tag.endswith("}p"):
                        yield ("paragraph", Paragraph(inner, doc), section_idx)
                    elif inner.tag.endswith("}tbl"):
                        yield ("table", Table(inner, doc), section_idx)


class ParagraphFormat:
    def __init__(self):
        self.left_indent = 0
        self.right_indent = 0
        self.first_line_indent = 0
        self.space_before = 0
        self.space_after = 0
        self.line_spacing = 16
        self.align = "left"
        self.bold = False
        self.italic = False
        self.size = 14
        self.font = "Inter"
        self.color = "#000000"
        self.has_border_bottom = False
        self.border_width = 2


def _reset_layout_defaults():
    global PAGE_WIDTH, PAGE_HEIGHT, LEFT_MARGIN, RIGHT_MARGIN, TOP_MARGIN, BOTTOM_MARGIN, CONTENT_WIDTH
    PAGE_WIDTH = DEFAULT_PAGE_WIDTH
    PAGE_HEIGHT = DEFAULT_PAGE_HEIGHT
    LEFT_MARGIN = DEFAULT_LEFT_MARGIN
    RIGHT_MARGIN = DEFAULT_RIGHT_MARGIN
    TOP_MARGIN = DEFAULT_TOP_MARGIN
    BOTTOM_MARGIN = DEFAULT_BOTTOM_MARGIN
    CONTENT_WIDTH = PAGE_WIDTH - (LEFT_MARGIN + RIGHT_MARGIN)


def _page_top(page_idx: int) -> int:
    return page_idx * (PAGE_HEIGHT + PAGE_GAP) + TOP_MARGIN


def _page_bottom_limit(page_idx: int, reserved: int = 40) -> int:
    return ((page_idx + 1) * PAGE_HEIGHT) + (page_idx * PAGE_GAP) - BOTTOM_MARGIN - reserved


def _extract_paragraph_text(par: Paragraph) -> str:
    """Извлекает текст абзаца, сохраняя переносы и табуляции."""
    chunks: list[str] = []
    for run in par.runs:
        if run.text:
            chunks.append(run.text)

        brs = run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br")
        for br in brs:
            if br.get(qn("w:type")) in ("textWrapping", None):
                chunks.append("\n")

        tabs = run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tab")
        for _ in tabs:
            chunks.append("\t")

    if not chunks:
        return par.text.strip()

    text = "".join(chunks).replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    normalized = "\n".join(lines).strip()
    return normalized


def _get_paragraph_format(par: Paragraph, style_cache: dict[str, ParagraphFormat]) -> ParagraphFormat:
    """Получение форматирования с учетом стилей"""
    fmt = ParagraphFormat()

    # === БАЗОВЫЙ СТИЛЬ ===
    if par.style and par.style.name:
        base_fmt = style_cache.get(par.style.name)
        if base_fmt:
            for attr in [
                "left_indent",
                "right_indent",
                "first_line_indent",
                "space_before",
                "space_after",
                "line_spacing",
                "align",
                "bold",
                "italic",
                "size",
                "font",
            ]:
                setattr(fmt, attr, getattr(base_fmt, attr))

    # === ФОРМАТИРОВАНИЕ ПАРАГРАФА ===
    try:
        par_fmt = par.paragraph_format
        if par_fmt.left_indent is not None:
            fmt.left_indent = int(par_fmt.left_indent.pt * PT_TO_PX)
        if par_fmt.right_indent is not None:
            fmt.right_indent = int(par_fmt.right_indent.pt * PT_TO_PX)
        if par_fmt.first_line_indent is not None:
            fmt.first_line_indent = int(par_fmt.first_line_indent.pt * PT_TO_PX)
        if par_fmt.space_before is not None:
            fmt.space_before = int(par_fmt.space_before.pt * PT_TO_PX)
        if par_fmt.space_after is not None:
            fmt.space_after = int(par_fmt.space_after.pt * PT_TO_PX)
        if par_fmt.line_spacing is not None:
            fmt.line_spacing = int(par_fmt.line_spacing * PT_TO_PX)

        if par.alignment is not None:
            align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
            fmt.align = align_map.get(par.alignment, "left")
    except Exception:
        pass

    # === RUNS ===
    _apply_runs_format(par, fmt)

    # === ГРАНИЦЫ ===
    _check_border_bottom(par, fmt)

    return fmt


def _apply_runs_format(par: Paragraph, fmt: ParagraphFormat):
    """Применение форматирования из runs"""
    if not par.runs:
        return

    main_run = next((r for r in par.runs if r.text.strip()), par.runs[0])

    if main_run.bold:
        fmt.bold = True
    if main_run.italic:
        fmt.italic = True

    if main_run.font:
        if main_run.font.size:
            fmt.size = int(main_run.font.size.pt)
        if main_run.font.name:
            fmt.font = main_run.font.name
        if main_run.font.color and main_run.font.color.rgb:
            fmt.color = f"#{main_run.font.color.rgb}"


def _check_border_bottom(par: Paragraph, fmt: ParagraphFormat):
    """Проверка наличия нижней границы"""
    try:
        pPr = par._element.pPr
        if pPr is None:
            return

        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            return

        bottom = pBdr.find(qn("w:bottom"))
        if bottom is not None and bottom.get(qn("w:val")) not in (None, "none", "nil"):
            fmt.has_border_bottom = True
            sz = bottom.get(qn("w:sz"))
            fmt.border_width = max(1, int((int(sz) or 4) * PT_TO_PX / 12))
    except Exception:
        pass


def _parse_table(table: Table, current_y: int, style_cache: dict[str, ParagraphFormat]) -> dict[str, Any] | None:
    """Парсинг таблицы с улучшенной обработкой данных"""
    rows_data = []
    row_heights = []
    cell_colors = []  # Массив цветов текста для каждой ячейки
    col_widths = _get_table_column_widths(table)

    # Сначала обрабатываем все ячейки
    for _row_idx, row in enumerate(table.rows):
        r_cells = []
        r_colors = []
        max_cell_height = 0

        for cell_idx, cell in enumerate(row.cells):
            # Извлекаем текст из всех параграфов ячейки
            cell_paragraphs = []
            cell_color = "#000000"  # Цвет по умолчанию

            for par in cell.paragraphs:
                par_text = _extract_paragraph_text(par)
                if par_text:
                    cell_paragraphs.append(par_text)
                    # Извлекаем цвет текста из первого параграфа с текстом
                    if cell_color == "#000000":
                        cell_color = _get_cell_text_color(par)

            # Объединяем параграфы
            cell_text = "\n".join(cell_paragraphs) if cell_paragraphs else ""

            # Очищаем текст от лишних пробелов, но сохраняем переносы строк
            # Заменяем множественные пробелы на один, но сохраняем \n
            lines = cell_text.split("\n")
            cell_text = "\n".join(" ".join(line.split()) for line in lines)

            r_cells.append(cell_text)
            r_colors.append(cell_color)

            # Вычисляем высоту ячейки
            cell_w = col_widths[cell_idx] if cell_idx < len(col_widths) else CONTENT_WIDTH // len(row.cells)
            cell_h = _calculate_table_cell_height(cell, cell_w, style_cache)

            if cell_h > max_cell_height:
                max_cell_height = cell_h

        rows_data.append(r_cells)
        cell_colors.append(r_colors)
        row_heights.append(max_cell_height)

    if not rows_data:
        return None

    # Обрабатываем объединенные ячейки
    merged_cells_map = _get_merged_cells_map(table)
    for (row_idx, cell_idx), (main_row, main_col) in merged_cells_map.items():
        if (
            main_row < len(rows_data)
            and main_col < len(rows_data[main_row])
            and row_idx < len(rows_data)
            and cell_idx < len(rows_data[row_idx])
        ):
            # Копируем данные из основной ячейки в объединенную
            rows_data[row_idx][cell_idx] = rows_data[main_row][main_col]
            if (
                main_row < len(cell_colors)
                and main_col < len(cell_colors[main_row])
                and row_idx < len(cell_colors)
                and cell_idx < len(cell_colors[row_idx])
            ):
                cell_colors[row_idx][cell_idx] = cell_colors[main_row][main_col]

    # Нормализуем количество колонок
    max_cols = max(len(r) for r in rows_data) if rows_data else 0
    for r in rows_data:
        while len(r) < max_cols:
            r.append("")
    for r in cell_colors:
        while len(r) < max_cols:
            r.append("#000000")

    if not rows_data:
        return None

    # Более консервативная высота: важнее не потерять строки, чем сделать плотнее.
    row_spacing = 8
    total_height = sum(row_heights) + (len(row_heights) * row_spacing) + 12

    # Минимум на строку под многострочный контент (особенно кириллица + переносы)
    min_table_height = len(rows_data) * 44
    total_height = max(total_height, min_table_height)

    return {
        "rows": rows_data,
        "row_heights": row_heights,
        "total_height": total_height,
        "cell_text_colors": cell_colors,
    }


def _calculate_table_cell_height(cell: _Cell, width_px: int, style_cache: dict[str, ParagraphFormat]) -> int:
    """ТОЧНЫЙ расчет высоты ячейки с улучшенной видимостью"""
    total_height = 16

    has_content = False
    for par in cell.paragraphs:
        cell_fmt = _get_paragraph_format(par, style_cache)
        text_content = par.text.strip()

        if text_content:
            has_content = True
            h = _calculate_text_height_exact(text_content, cell_fmt.size, width_px - 16, cell_fmt.line_spacing)
            total_height += h + cell_fmt.space_before + cell_fmt.space_after
        else:
            # Для пустых параграфов используем минимальную высоту
            min_par_height = max(24, int(cell_fmt.size * 1.7))
            total_height += min_par_height

    # Минимальная высота ячейки для лучшей видимости
    min_cell_height = 42
    if not has_content:
        # Если ячейка пустая, все равно даем минимальную высоту
        total_height = min_cell_height

    total_height = max(total_height, min_cell_height)
    total_height += 8

    return total_height


def _calculate_text_height_exact(text: str, font_size: int, width_px: int, line_spacing: float) -> int:
    """ТОЧНЕЙШАЯ высота текста с переносами"""
    if not text:
        return max(20, int(font_size * 1.5))  # Минимальная высота даже для пустого текста

    # Консервативная оценка: лучше дать больше высоты, чем обрезать строки.
    base_line_h = max(int(font_size * 1.8), int(line_spacing * 1.15)) if line_spacing > 0 else int(font_size * 1.8)
    char_w = font_size * 0.72

    # Символов в строке
    chars_per_line = max(1, int(width_px // char_w))

    total_lines = 0
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            total_lines += 1
            continue

        words = line.split()
        current_line_len = 0

        for word in words:
            word_len = len(word)
            if word_len > chars_per_line:
                # Длинное слово переносится
                wrapped_lines = math.ceil(word_len / chars_per_line)
                total_lines += wrapped_lines
                current_line_len = word_len % chars_per_line
            elif current_line_len + word_len + 1 <= chars_per_line:
                current_line_len += word_len + 1
            else:
                total_lines += 1
                current_line_len = word_len + 1

        total_lines += 1

    # Добавляем padding сверху и снизу для лучшей видимости
    calculated_height = total_lines * base_line_h
    min_height = max(30, int(font_size * 2.2))
    return max(calculated_height, min_height) + 12


def _get_table_column_widths(table: Table) -> list[int]:
    """ТОЧНЫЕ ширины колонок из XML"""
    try:
        tbl_grid = table._element.find(qn("w:tblGrid"))
        if tbl_grid is not None:
            widths = []
            for grid_col in tbl_grid.findall(qn("w:gridCol")):
                w = grid_col.get(qn("w:w"))
                if w:
                    # twip to px: 1 twip = 1/20 pt = 0.0666667 px
                    width_px = int(int(w) * 0.0666667)
                    widths.append(width_px)
            if widths:
                return widths

        # Резервный вариант
        if table.rows:
            cell_count = len(table.rows[0].cells)
            if cell_count > 0:
                return [CONTENT_WIDTH // cell_count] * cell_count

    except Exception:
        pass

    return [CONTENT_WIDTH]


def _has_page_break(par: Paragraph) -> bool:
    """Проверка разрыва страницы"""
    try:
        pPr = par._element.pPr
        if pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None:
            return True

        for run in par.runs:
            brs = run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br")
            for br in brs:
                if br.get(qn("w:type")) == "page":
                    return True
    except Exception:
        pass
    return False


def _load_document_margins(doc: DocumentObject):
    """Загрузка полей документа"""
    global PAGE_WIDTH, PAGE_HEIGHT, LEFT_MARGIN, RIGHT_MARGIN, TOP_MARGIN, BOTTOM_MARGIN, CONTENT_WIDTH

    try:
        if doc.sections:
            section = doc.sections[0]
            if section.page_width:
                PAGE_WIDTH = int(section.page_width.pt * PT_TO_PX)
            if section.page_height:
                PAGE_HEIGHT = int(section.page_height.pt * PT_TO_PX)
            if section.left_margin:
                LEFT_MARGIN = int(section.left_margin.pt * PT_TO_PX)
            if section.right_margin:
                RIGHT_MARGIN = int(section.right_margin.pt * PT_TO_PX)
            if section.top_margin:
                TOP_MARGIN = int(section.top_margin.pt * PT_TO_PX)
            if section.bottom_margin:
                BOTTOM_MARGIN = int(section.bottom_margin.pt * PT_TO_PX)
            CONTENT_WIDTH = PAGE_WIDTH - (LEFT_MARGIN + RIGHT_MARGIN)
    except Exception:
        pass


def _get_table_format(table: Table) -> dict[str, int]:
    """Получение форматирования таблицы (отступы, выравнивание)"""
    format_dict = {"left_indent": 0, "right_indent": 0, "space_before": 8, "space_after": 12}

    try:
        tbl_pr = table._element.tblPr
        if tbl_pr is not None:
            # Отступ слева
            tbl_ind = tbl_pr.find(qn("w:tblInd"))
            if tbl_ind is not None:
                w_val = tbl_ind.get(qn("w:w"))
                if w_val:
                    format_dict["left_indent"] = int(int(w_val) * 0.0666667)  # twip to px

            # Отступ справа (через tblCellMar)
            tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
            if tbl_cell_mar is not None:
                left_mar = tbl_cell_mar.find(qn("w:left"))
                right_mar = tbl_cell_mar.find(qn("w:right"))
                if left_mar is not None:
                    w_val = left_mar.get(qn("w:w"))
                    if w_val:
                        format_dict["left_indent"] += int(int(w_val) * 0.0666667)
                if right_mar is not None:
                    w_val = right_mar.get(qn("w:w"))
                    if w_val:
                        format_dict["right_indent"] = int(int(w_val) * 0.0666667)
    except Exception:
        pass

    return format_dict


def _get_cell_text_color(par: Paragraph) -> str:
    """Извлечение цвета текста из параграфа ячейки"""
    try:
        # Проверяем цвет в runs
        for run in par.runs:
            if run.font and run.font.color and run.font.color.rgb:
                return f"#{run.font.color.rgb}"

        # Проверяем цвет в стиле параграфа
        if par.style and par.style.font and par.style.font.color and par.style.font.color.rgb:
            return f"#{par.style.font.color.rgb}"
    except Exception:
        pass

    return "#000000"  # Цвет по умолчанию


def _get_merged_cells_map(table: Table) -> dict[tuple, tuple]:
    """Получение карты объединенных ячеек (row, col) -> (main_row, main_col)

    В python-docx, когда ячейка объединена вертикально (vMerge),
    она все равно присутствует в row.cells, но может быть пустой.
    Для горизонтальных объединений (gridSpan), ячейка занимает несколько колонок,
    но в row.cells она представлена как одна ячейка.
    """
    merged_map = {}

    try:
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                tc_pr = cell._element.find(qn("w:tcPr"))
                if tc_pr is None:
                    continue

                v_merge = tc_pr.find(qn("w:vMerge"))
                if v_merge is None:
                    continue

                v_val = v_merge.get(qn("w:val"))
                if v_val in (None, "continue"):
                    for prev_row_idx in range(row_idx - 1, -1, -1):
                        if cell_idx >= len(table.rows[prev_row_idx].cells):
                            continue

                        prev_tc_pr = table.rows[prev_row_idx].cells[cell_idx]._element.find(qn("w:tcPr"))
                        if prev_tc_pr is None:
                            continue
                        prev_v_merge = prev_tc_pr.find(qn("w:vMerge"))
                        if prev_v_merge is None:
                            continue

                        prev_val = prev_v_merge.get(qn("w:val"))
                        if prev_val in ("restart", None):
                            merged_map[(row_idx, cell_idx)] = (prev_row_idx, cell_idx)
                            break
    except Exception as e:
        print(f"Error processing merged cells: {e}")

    return merged_map
